"""
Engine9: DocxFinalTranslatorEngine + Custom Dictionary
=======================================================
Upgrade dari Engine 8 dengan fitur KAMUS KHUSUS (Custom Dictionary):

  - Kamus bisa diisi manual (add_term)
  - Kamus bisa diimpor dari file CSV
  - Kamus bisa diimpor dari file Excel (.xlsx/.xls)
  - Kamus bawaan (built-in) untuk istilah teknis umum
  - Istilah kamus DIPRIORITASKAN → tidak dikirim ke Google Translate
    (Teknik: term → placeholder token → translate → restore target term)

Cara pakai CustomDictionary:
    d = CustomDictionary()
    d.load_defaults()                         # muat istilah bawaan
    d.load_from_csv("kamus.csv")              # CSV: kolom source,target
    d.load_from_excel("kamus.xlsx", sheet="Sheet1", src_col="Inggris", tgt_col="Indonesia")
    d.add_term("compressive strength", "kuat tekan")

    engine = DocxFinalTranslatorEngine(custom_dict=d)
    engine.translate("input.docx", "output.docx")

Format CSV minimal:
    source,target
    compressive strength,kuat tekan
    rebar,tulangan baja

Format Excel minimal:
    Kolom A (header "source" atau nama custom) → istilah asing
    Kolom B (header "target" atau nama custom) → terjemahan Indonesia

Dependensi:
    pip install deep-translator python-docx lxml openpyxl pandas
"""

import re
import copy
import time
import uuid
import traceback
import csv
import os

from docx import Document
from docx.oxml.ns import qn
import lxml.etree as etree


# ─────────────────────────────────────────────────────────────────────────────
# NAMESPACE
# ─────────────────────────────────────────────────────────────────────────────

_NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_W    = f'{{{_NS_W}}}'


# ─────────────────────────────────────────────────────────────────────────────
# KONSTANTA
# ─────────────────────────────────────────────────────────────────────────────

_RE_PURE_NUMBER = re.compile(
    r'^[\d\s\.\,\:\;\-\(\)\[\]\/\\\+\=\*\%\&\^\$\#\@\!\"\'`~<>{}|_]+$'
)
_RE_COPYRIGHT = re.compile(r'©|BSN\s*\d{4}', re.IGNORECASE)

_SKIP_STYLES = {
    'caption', 'header', 'footer',
    'toc 1', 'toc 2', 'toc 3', 'toc 4', 'toc 5',
    'table of figures', 'footnote text', 'endnote text', 'macro text',
}
_BIBLIO_TITLE_STYLES = {'biblioti', 'bibliotitle', 'bibliography title'}
_BIBLIO_KEYWORDS_EXACT = {
    'bibliografi', 'bibliography',
    'daftar acuan', 'daftar pustaka', 'daftar referensi',
}
_ANNEX_STYLE_IDS = {'ANNEX', 'Annex', 'annex'}
_HEADING_STYLES_WITH_NUM = {
    'Heading1', 'Heading2', 'Heading3',
    'ANNEX', 'a2', 'a3',
    'Heading4', 'Heading5', 'Heading6',
}
_TRANSLATE_DELAY = 0.15
_EM_DASH = '—'


# ─────────────────────────────────────────────────────────────────────────────
# CUSTOM DICTIONARY
# ─────────────────────────────────────────────────────────────────────────────

# Istilah bawaan — bidang konstruksi / standar teknik
_DEFAULT_TERMS = {
    # Umum Teknik Sipil & Struktur
    "compressive strength"          : "kuat tekan",
    "tensile strength"              : "kuat tarik",
    "flexural strength"             : "kuat lentur",
    "shear strength"                : "kuat geser",
    "bearing capacity"              : "daya dukung",
    "allowable stress"              : "tegangan izin",
    "ultimate load"                 : "beban ultimit",
    "serviceability"                : "kemampuan layan",
    "reinforced concrete"           : "beton bertulang",
    "prestressed concrete"          : "beton prategang",
    "reinforcement bar"             : "tulangan baja",
    "rebar"                         : "tulangan baja",
    "stirrup"                       : "sengkang",
    "column"                        : "kolom",
    "beam"                          : "balok",
    "slab"                          : "pelat",
    "footing"                       : "pondasi tapak",
    "pile"                          : "tiang pancang",
    "shear wall"                    : "dinding geser",
    "moment of inertia"             : "momen inersia",
    "modulus of elasticity"         : "modulus elastisitas",
    "yield strength"                : "kuat leleh",
    "nominal strength"              : "kuat nominal",
    "design strength"               : "kuat rencana",
    "load factor"                   : "faktor beban",
    "resistance factor"             : "faktor ketahanan",
    "live load"                     : "beban hidup",
    "dead load"                     : "beban mati",
    "wind load"                     : "beban angin",
    "seismic load"                  : "beban gempa",
    "earthquake"                    : "gempa bumi",
    "lateral force"                 : "gaya lateral",
    "drift"                         : "simpangan",
    "displacement"                  : "perpindahan",
    "deflection"                    : "lendutan",
    "crack width"                   : "lebar retak",
    "cover"                         : "selimut beton",
    "splice"                        : "sambungan",
    "development length"            : "panjang penyaluran",
    "embedment length"              : "panjang penanaman",
    "anchorage"                     : "angkur",

    # Material
    "aggregate"                     : "agregat",
    "coarse aggregate"              : "agregat kasar",
    "fine aggregate"                : "agregat halus",
    "cement"                        : "semen",
    "water-cement ratio"            : "rasio air-semen",
    "admixture"                     : "bahan tambah",
    "curing"                        : "perawatan",
    "mix design"                    : "rancang campur",
    "slump"                         : "slump",
    "workability"                   : "kemudahan pengerjaan",

    # Istilah Standar / Dokumen
    "annex"                         : "lampiran",
    "clause"                        : "pasal",
    "subclause"                     : "subpasal",
    "commentary"                    : "komentar",
    "informative"                   : "informatif",
    "normative"                     : "normatif",
    "shall"                         : "harus",
    "should"                        : "sebaiknya",
    "may"                           : "boleh",
    "table"                         : "tabel",
    "figure"                        : "gambar",
    "equation"                      : "persamaan",
    "notation"                      : "notasi",
    "symbol"                        : "simbol",
    "bibliography"                  : "bibliografi",
    "reference"                     : "acuan",
    "standard"                      : "standar",

    # Lain-lain Umum
    "quality control"               : "pengendalian mutu",
    "quality assurance"             : "jaminan mutu",
    "inspection"                    : "inspeksi",
    "testing"                       : "pengujian",
    "specification"                 : "spesifikasi",
    "tolerance"                     : "toleransi",
    "safety factor"                 : "faktor keamanan",
    "factor of safety"              : "faktor keamanan",
}


class CustomDictionary:
    """
    Kamus istilah khusus yang digunakan SEBELUM penerjemahan Google.

    Cara kerja:
      1. Setiap istilah sumber (source) di dalam teks digantikan sementara
         dengan sebuah TOKEN unik (misal: @@TK_0001@@) agar Google Translate
         tidak mengacau terjemahan istilah tersebut.
      2. Setelah terjemahan selesai, TOKEN dikembalikan menjadi
         terjemahan target yang sudah ditentukan di kamus.

    Entries: {source_lower → (source_original, target)}
    """

    def __init__(self):
        # key: lowercase source → (original_source, target)
        self._entries: dict[str, tuple[str, str]] = {}

    # ── TAMBAH MANUAL ──────────────────────────────────────────────────────

    def add_term(self, source: str, target: str) -> None:
        """Tambah satu istilah secara manual."""
        s = source.strip()
        t = target.strip()
        if s and t:
            self._entries[s.lower()] = (s, t)

    def remove_term(self, source: str) -> None:
        """Hapus istilah dari kamus."""
        self._entries.pop(source.strip().lower(), None)

    def clear(self) -> None:
        """Kosongkan semua entri."""
        self._entries.clear()

    # ── KAMUS BAWAAN ───────────────────────────────────────────────────────

    def load_defaults(self) -> int:
        """
        Muat istilah bawaan (bidang konstruksi & standar teknik).
        Returns jumlah istilah yang ditambahkan.
        """
        before = len(self._entries)
        for src, tgt in _DEFAULT_TERMS.items():
            self._entries[src.lower()] = (src, tgt)
        return len(self._entries) - before

    # ── DARI CSV ───────────────────────────────────────────────────────────

    def load_from_csv(
        self,
        filepath: str,
        src_col:  str = 'source',
        tgt_col:  str = 'target',
        delimiter: str = ',',
        encoding:  str = 'utf-8-sig',
    ) -> int:
        """
        Muat kamus dari file CSV.

        Format yang didukung:
          - Ada header: kolom bernama src_col dan tgt_col
          - Tanpa header: kolom pertama = source, kolom kedua = target

        Returns jumlah istilah yang berhasil ditambahkan.
        """
        if not os.path.isfile(filepath):
            raise FileNotFoundError(f"File CSV tidak ditemukan: {filepath}")

        count = 0
        with open(filepath, newline='', encoding=encoding) as f:
            sample = f.read(1024)
            f.seek(0)
            has_header = csv.Sniffer().has_header(sample)

            reader = csv.DictReader(f, delimiter=delimiter) if has_header else \
                     csv.reader(f, delimiter=delimiter)

            for row in reader:
                if has_header:
                    src = row.get(src_col, row.get('source', '')).strip()
                    tgt = row.get(tgt_col, row.get('target', '')).strip()
                else:
                    row = list(row)
                    if len(row) < 2:
                        continue
                    src, tgt = row[0].strip(), row[1].strip()

                if src and tgt:
                    self._entries[src.lower()] = (src, tgt)
                    count += 1

        return count

    # ── DARI EXCEL ─────────────────────────────────────────────────────────

    def load_from_excel(
        self,
        filepath:   str,
        sheet_name: str | int = 0,
        src_col:    str = 'source',
        tgt_col:    str = 'target',
    ) -> int:
        """
        Muat kamus dari file Excel (.xlsx / .xls).

        Aturan kolom:
          - Jika header row berisi src_col dan tgt_col, gunakan nama itu.
          - Jika tidak ditemukan, kolom pertama = source, kolom kedua = target.

        Returns jumlah istilah yang berhasil ditambahkan.
        """
        if not os.path.isfile(filepath):
            raise FileNotFoundError(f"File Excel tidak ditemukan: {filepath}")

        try:
            import pandas as pd
        except ImportError:
            raise ImportError("Jalankan: pip install pandas openpyxl")

        df = pd.read_excel(filepath, sheet_name=sheet_name, dtype=str)
        df = df.fillna('')

        # Cari kolom source
        col_src = _find_col(df.columns.tolist(), [src_col, 'source', 'Source', 'SOURCE',
                                                    'Inggris', 'inggris', 'Bahasa Asing',
                                                    'asing', 'English', 'english'])
        # Cari kolom target
        col_tgt = _find_col(df.columns.tolist(), [tgt_col, 'target', 'Target', 'TARGET',
                                                    'Indonesia', 'indonesia', 'Bahasa Indonesia',
                                                    'id', 'ID'])

        # Fallback ke kolom pertama & kedua
        if col_src is None:
            col_src = df.columns[0]
        if col_tgt is None and len(df.columns) >= 2:
            col_tgt = df.columns[1]
        if col_tgt is None:
            raise ValueError("Kolom target tidak ditemukan di Excel.")

        count = 0
        for _, row in df.iterrows():
            src = str(row[col_src]).strip()
            tgt = str(row[col_tgt]).strip()
            if src and tgt and src.lower() not in ('nan', '') and tgt.lower() not in ('nan', ''):
                self._entries[src.lower()] = (src, tgt)
                count += 1

        return count

    # ── DARI GOOGLE SHEET ──────────────────────────────────────────────────

    def load_from_google_sheet(
        self,
        url:      str,
        src_col:  str = 'source',
        tgt_col:  str = 'target',
        timeout:  int = 15,
    ) -> int:
        """
        Muat kamus dari Google Sheets melalui URL shareable.

        Google Sheet harus diset "Anyone with the link can view".

        URL yang didukung (semua format otomatis dikonversi ke CSV export):
          - https://docs.google.com/spreadsheets/d/SHEET_ID/edit?usp=sharing
          - https://docs.google.com/spreadsheets/d/SHEET_ID/edit#gid=GID
          - https://docs.google.com/spreadsheets/d/SHEET_ID/pub...
          - URL CSV export langsung

        Format sheet:
          Baris pertama = header. Kolom source (istilah asing) dan target
          (terjemahan Indonesia). Nama kolom fleksibel — bisa source/target,
          Inggris/Indonesia, asing/id, dll.

        Returns jumlah istilah yang berhasil dimuat.
        """
        try:
            import urllib.request
            import io
        except ImportError:
            raise ImportError("urllib adalah bagian dari Python standar, seharusnya tersedia.")

        csv_url = _google_sheet_to_csv_url(url)

        try:
            req = urllib.request.Request(
                csv_url,
                headers={'User-Agent': 'Mozilla/5.0'}
            )
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                raw = resp.read().decode('utf-8-sig')
        except Exception as e:
            raise ConnectionError(
                f"Gagal mengambil data dari Google Sheets.\n"
                f"Pastikan sheet diset 'Anyone with the link can view'.\n"
                f"Error: {e}"
            )

        # Parse sebagai CSV
        f = io.StringIO(raw)
        reader = csv.DictReader(f)
        fieldnames = reader.fieldnames or []

        # Cari kolom source & target
        col_src = _find_col(fieldnames, [src_col, 'source', 'Source', 'SOURCE',
                                          'Inggris', 'inggris', 'Bahasa Asing',
                                          'asing', 'English', 'english'])
        col_tgt = _find_col(fieldnames, [tgt_col, 'target', 'Target', 'TARGET',
                                          'Indonesia', 'indonesia', 'Bahasa Indonesia',
                                          'id', 'ID'])

        # Fallback ke kolom pertama & kedua jika nama tidak cocok
        if col_src is None and len(fieldnames) >= 1:
            col_src = fieldnames[0]
        if col_tgt is None and len(fieldnames) >= 2:
            col_tgt = fieldnames[1]
        if col_tgt is None:
            raise ValueError(
                f"Kolom target tidak ditemukan. "
                f"Header yang tersedia: {fieldnames}. "
                f"Gunakan nama kolom: source, target, Inggris, Indonesia, dll."
            )

        count = 0
        for row in reader:
            src = str(row.get(col_src, '')).strip()
            tgt = str(row.get(col_tgt, '')).strip()
            if src and tgt and src.lower() not in ('', 'nan') and tgt.lower() not in ('', 'nan'):
                self._entries[src.lower()] = (src, tgt)
                count += 1

        return count

    # ── SIMPAN KE CSV ──────────────────────────────────────────────────────

    def save_to_csv(self, filepath: str, encoding: str = 'utf-8-sig') -> int:
        """Simpan semua entri kamus ke file CSV."""
        count = 0
        with open(filepath, 'w', newline='', encoding=encoding) as f:
            writer = csv.writer(f)
            writer.writerow(['source', 'target'])
            for src_lower, (src_orig, tgt) in sorted(self._entries.items()):
                writer.writerow([src_orig, tgt])
                count += 1
        return count

    # ── SIMPAN KE EXCEL ────────────────────────────────────────────────────

    def save_to_excel(self, filepath: str, sheet_name: str = 'Kamus') -> int:
        """Simpan semua entri kamus ke file Excel."""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("Jalankan: pip install pandas openpyxl")

        rows = [{'source': src_orig, 'target': tgt}
                for _, (src_orig, tgt) in sorted(self._entries.items())]
        df = pd.DataFrame(rows, columns=['source', 'target'])
        df.to_excel(filepath, sheet_name=sheet_name, index=False)
        return len(rows)

    # ── INFO ───────────────────────────────────────────────────────────────

    def __len__(self) -> int:
        return len(self._entries)

    def __repr__(self) -> str:
        return f"CustomDictionary({len(self._entries)} terms)"

    def list_terms(self) -> list[tuple[str, str]]:
        """Kembalikan daftar (source, target) terurut."""
        return [(src, tgt) for _, (src, tgt) in sorted(self._entries.items())]

    # ── INTERNAL: APPLY BEFORE/AFTER TRANSLATE ────────────────────────────

    def _apply_pre(self, text: str) -> tuple[str, dict]:
        """
        Ganti istilah kamus dengan placeholder token.
        Returns (text_with_tokens, {token: target}).
        Token bersifat unik dan tidak akan diterjemahkan Google.
        """
        if not self._entries:
            return text, {}

        token_map = {}  # token → target
        result    = text

        # Urutkan dari yang terpanjang ke terpendek agar frasa multi-kata
        # dicocokkan sebelum kata tunggal
        sorted_entries = sorted(
            self._entries.items(),
            key=lambda x: len(x[0]),
            reverse=True,
        )

        for src_lower, (src_orig, tgt) in sorted_entries:
            # Cari secara case-insensitive, whole-word
            pattern = re.compile(
                r'(?<![A-Za-z0-9])' + re.escape(src_lower) + r'(?![A-Za-z0-9])',
                re.IGNORECASE,
            )
            if pattern.search(result):
                token = f'@@TK_{uuid.uuid4().hex[:8].upper()}@@'
                token_map[token] = tgt
                result = pattern.sub(token, result)

        return result, token_map

    def _apply_post(self, translated: str, token_map: dict) -> str:
        """
        Kembalikan token placeholder menjadi terjemahan target.
        """
        result = translated
        for token, tgt in token_map.items():
            # Google Translate kadang menambah spasi atau huruf kapital di token
            # Coba pencocokan fleksibel
            result = re.sub(re.escape(token), tgt, result, flags=re.IGNORECASE)
        return result


def _find_col(columns: list, candidates: list) -> str | None:
    """Cari nama kolom pertama yang cocok dari daftar kandidat."""
    for c in candidates:
        if c in columns:
            return c
    return None


def _google_sheet_to_csv_url(url: str) -> str:
    """
    Konversi berbagai format URL Google Sheets ke URL CSV export.

    Contoh input  → output:
      .../spreadsheets/d/SHEET_ID/edit?usp=sharing
        → .../spreadsheets/d/SHEET_ID/export?format=csv

      .../spreadsheets/d/SHEET_ID/edit#gid=123456
        → .../spreadsheets/d/SHEET_ID/export?format=csv&gid=123456

      .../spreadsheets/d/SHEET_ID/pub?output=csv   (sudah CSV)
        → dikembalikan apa adanya

      URL CSV langsung (bukan spreadsheets Google)
        → dikembalikan apa adanya
    """
    url = url.strip()

    # Sudah berupa CSV export atau bukan Google Sheets → return apa adanya
    if 'output=csv' in url or 'format=csv' in url:
        return url
    if 'docs.google.com/spreadsheets' not in url:
        return url

    # Ekstrak SHEET_ID
    m = re.search(r'/spreadsheets/d/([a-zA-Z0-9_-]+)', url)
    if not m:
        raise ValueError(f"Tidak dapat mengekstrak Sheet ID dari URL: {url}")
    sheet_id = m.group(1)

    # Ekstrak gid (tab/sheet tertentu) jika ada
    gid_match = re.search(r'[#&?]gid=(\d+)', url)
    gid_param = f'&gid={gid_match.group(1)}' if gid_match else ''

    return f'https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv{gid_param}'


# ─────────────────────────────────────────────────────────────────────────────
# HELPER: FILTER
# ─────────────────────────────────────────────────────────────────────────────

def _skip_text(text: str) -> bool:
    t = text.strip()
    if len(t) < 3:                   return True
    if _RE_PURE_NUMBER.fullmatch(t): return True
    if _RE_COPYRIGHT.search(t):      return True
    return False

def _skip_paragraph(para, past_bibliography: bool = False) -> bool:
    if past_bibliography:
        return True
    if not para.text.strip():
        return True
    for tag in [f'{_W}drawing', f'{_W}pict']:
        if para._element.find('.//' + tag) is not None:
            return True
    style_name = (para.style.name or '').lower()
    if any(style_name.startswith(s) for s in _SKIP_STYLES):
        return True
    return False

def _is_biblio_title_para(para) -> bool:
    style_id = ''
    try:
        pStyle = para._element.find(f'{_W}pPr/{_W}pStyle')
        if pStyle is not None:
            style_id = pStyle.get(f'{_W}val', '').lower()
    except Exception:
        pass
    if style_id in _BIBLIO_TITLE_STYLES:
        return True
    txt = para.text.strip().lower()
    return bool(txt) and not txt[0].isdigit() and txt in _BIBLIO_KEYWORDS_EXACT

def _get_para_style_id(para) -> str:
    pStyle = para._element.find(f'{_W}pPr/{_W}pStyle')
    if pStyle is not None:
        return pStyle.get(f'{_W}val', '')
    return ''

def _get_style_id(el) -> str:
    pStyle = el.find(f'{_W}pPr/{_W}pStyle')
    return pStyle.get(f'{_W}val', '') if pStyle is not None else 'Normal'


# ─────────────────────────────────────────────────────────────────────────────
# FITUR 1: REKONSTRUKSI ANNEX
# ─────────────────────────────────────────────────────────────────────────────

def _fix_annex_style_para(para) -> None:
    sid = _get_para_style_id(para)
    if sid not in _ANNEX_STYLE_IDS:
        return

    full_text = para.text.strip()
    if not full_text:
        return

    tag_norm = None
    title_part = full_text

    tags_to_find = [
        '(informatif)', '(normatif)',
        '(informative)', '(normative)',
        '(informasi)'
    ]

    for t in tags_to_find:
        idx = full_text.lower().find(t)
        if idx != -1:
            if 'norm' in t.lower():
                tag_norm = '(normatif)'
            else:
                tag_norm = '(informatif)'
            part_before = full_text[:idx]
            part_after  = full_text[idx + len(t):]
            title_part  = part_before + " " + part_after
            break

    title_part = re.sub(r'^Annex\s+[A-Z0-9\.]+\s*', '', title_part, flags=re.IGNORECASE).strip()

    pPr = para._element.find(f'{_W}pPr')
    for child in list(para._element):
        if child is not pPr:
            para._element.remove(child)

    def make_arial_run(text, is_bold=False):
        esc   = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        b_tag = '<w:b/><w:bCs/>' if is_bold else ''
        return etree.fromstring(
            f'<w:r xmlns:w="{_NS_W}">'
            f'<w:rPr>'
            f'<w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
            f'{b_tag}'
            f'<w:sz w:val="22"/><w:szCs w:val="22"/>'
            f'</w:rPr>'
            f'<w:t xml:space="preserve">{esc}</w:t>'
            f'</w:r>'
        )

    def make_br_run():
        return etree.fromstring(f'<w:r xmlns:w="{_NS_W}"><w:br/></w:r>')

    new_runs = [make_br_run()]
    if tag_norm:
        new_runs.append(make_arial_run(tag_norm, is_bold=False))
    if title_part:
        new_runs.append(make_br_run())
        new_runs.append(make_br_run())
        new_runs.append(make_arial_run(title_part, is_bold=True))

    for run_el in new_runs:
        para._element.append(run_el)


# ─────────────────────────────────────────────────────────────────────────────
# FITUR 2: EM DASH TO BULLETS
# ─────────────────────────────────────────────────────────────────────────────

def _get_or_create_emdash_numid(doc: Document) -> str:
    try:
        np = doc.part.numbering_part
        if np is None: return None
        nxml = np._element

        target_abstract_id = None
        for ab in nxml.findall(f'{_W}abstractNum'):
            for lvl in ab.findall(f'{_W}lvl'):
                txt_el = lvl.find(f'{_W}lvlText')
                if txt_el is not None and txt_el.get(f'{_W}val') == _EM_DASH:
                    target_abstract_id = ab.get(f'{_W}abstractNumId')
                    break
            if target_abstract_id: break

        if target_abstract_id:
            existing_nums = nxml.findall(f'{_W}num')
            max_num_id    = max((int(n.get(f'{_W}numId', 0)) for n in existing_nums), default=0)
            new_num_id    = str(max_num_id + 1)
            new_num       = etree.fromstring(
                f'<w:num xmlns:w="{_NS_W}" w:numId="{new_num_id}">'
                f'<w:abstractNumId w:val="{target_abstract_id}"/>'
                f'</w:num>'
            )
            nxml.append(new_num)
            return new_num_id

        existing_abstracts = nxml.findall(f'{_W}abstractNum')
        max_abstract_id    = max((int(a.get(f'{_W}abstractNumId', 0)) for a in existing_abstracts), default=0)
        new_abstract_id    = str(max_abstract_id + 1)
        existing_nums      = nxml.findall(f'{_W}num')
        max_num_id         = max((int(n.get(f'{_W}numId', 0)) for n in existing_nums), default=0)
        new_num_id         = str(max_num_id + 1)

        abstract_xml = f'''
        <w:abstractNum xmlns:w="{_NS_W}" w:abstractNumId="{new_abstract_id}">
            <w:multiLevelType w:val="hybridMultilevel"/>
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="bullet"/>
                <w:lvlText w:val="{_EM_DASH}"/>
                <w:lvlJc w:val="left"/>
                <w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr>
                <w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/></w:rPr>
            </w:lvl>
        </w:abstractNum>
        '''
        nxml.insert(0, etree.fromstring(abstract_xml))

        num_xml = f'''
        <w:num xmlns:w="{_NS_W}" w:numId="{new_num_id}">
            <w:abstractNumId w:val="{new_abstract_id}"/>
        </w:num>
        '''
        nxml.append(etree.fromstring(num_xml))
        return new_num_id

    except Exception as e:
        print(f"[Engine9] Error numbering: {e}")
        return None


def _convert_emdash_to_bullets(doc: Document) -> None:
    num_id = _get_or_create_emdash_numid(doc)
    if not num_id: return

    for para in doc.paragraphs:
        sid = _get_para_style_id(para)
        if sid in _ANNEX_STYLE_IDS or sid.startswith('Heading'):
            continue
        text = para.text.strip()
        if text.startswith(_EM_DASH):
            for r in para.runs:
                if _EM_DASH in r.text:
                    r.text = r.text.replace(_EM_DASH, "", 1).lstrip()
                    break

            pPr = para._element.find(f'{_W}pPr')
            if pPr is None:
                pPr = etree.SubElement(para._element, f'{_W}pPr')

            old_num = pPr.find(f'{_W}numPr')
            if old_num is not None: pPr.remove(old_num)

            pPr.insert(0, etree.fromstring(
                f'<w:numPr xmlns:w="{_NS_W}">'
                f'<w:ilvl w:val="0"/>'
                f'<w:numId w:val="{num_id}"/>'
                f'</w:numPr>'
            ))

            old_ind = pPr.find(f'{_W}ind')
            if old_ind is not None: pPr.remove(old_ind)
            pPr.insert(1, etree.fromstring(
                f'<w:ind xmlns:w="{_NS_W}" w:left="360" w:hanging="360"/>'
            ))


# ─────────────────────────────────────────────────────────────────────────────
# HELPER: ITALIC & SECTION BREAK
# ─────────────────────────────────────────────────────────────────────────────

def _has_inline_sectpr(para) -> bool:
    pPr = para._element.find(f'{_W}pPr')
    if pPr is None: return False
    return pPr.find(f'{_W}sectPr') is not None

def _all_runs_italic(para) -> bool:
    text_runs = [r for r in para.runs if r.text.strip()]
    if not text_runs: return False
    para_style_italic = False
    try:
        if para.style and para.style.font and para.style.font.italic:
            para_style_italic = True
    except Exception: pass
    for run in text_runs:
        italic = False
        if run.font.italic is True: italic = True
        if not italic:
            rPr = run._element.find(f'{_W}rPr')
            if rPr is not None:
                i_el = rPr.find(f'{_W}i')
                if i_el is not None:
                    val = i_el.get(f'{_W}val', 'true')
                    if val.lower() not in ('false', '0'): italic = True
        if not italic and para_style_italic:
            rPr = run._element.find(f'{_W}rPr')
            if rPr is not None:
                i_el = rPr.find(f'{_W}i')
                if i_el is not None:
                    val = i_el.get(f'{_W}val', 'true')
                    if val.lower() not in ('false', '0'): italic = True
                else: italic = True
            else: italic = True
        if not italic: return False
    return True


# ─────────────────────────────────────────────────────────────────────────────
# HELPER: BIBLIOGRAFI & AUTONUMBERING
# ─────────────────────────────────────────────────────────────────────────────

def _el_text(el) -> str:
    return ''.join(t.text or '' for t in el.findall(f'.//{_W}t')).strip()

def _is_bibliography_el(el) -> bool:
    if el.tag != f'{_W}p': return False
    sid = _get_style_id(el).lower()
    if sid in _BIBLIO_TITLE_STYLES: return True
    text = _el_text(el).lower().strip()
    if not text or text[0].isdigit(): return False
    return text in _BIBLIO_KEYWORDS_EXACT

def _find_bib_index(body_els: list) -> int:
    for i, el in enumerate(body_els):
        if _is_bibliography_el(el): return i
    return -1

def _find_content_start_before_bib(body_els: list, bib_idx: int) -> int:
    search_limit = bib_idx if bib_idx >= 0 else len(body_els)
    last_sectpr  = -1
    for i, el in enumerate(body_els):
        if i >= search_limit: break
        if el.tag != f'{_W}p': continue
        pPr = el.find(f'{_W}pPr')
        if pPr is not None and pPr.find(f'{_W}sectPr') is not None:
            last_sectpr = i
    return last_sectpr + 1

def _read_style_numpr(doc: Document) -> dict:
    result = {}
    try:
        styles_xml = doc.part.styles._element
        for style_el in styles_xml.findall(f'{_W}style'):
            sid = style_el.get(f'{_W}styleId', '')
            if sid not in _HEADING_STYLES_WITH_NUM: continue
            pPr = style_el.find(f'{_W}pPr')
            if pPr is None: continue
            numPr = pPr.find(f'{_W}numPr')
            if numPr is None: continue
            numId_el = numPr.find(f'{_W}numId')
            ilvl_el  = numPr.find(f'{_W}ilvl')
            nid  = numId_el.get(f'{_W}val', '0') if numId_el is not None else '0'
            ilvl = ilvl_el.get(f'{_W}val',  '0') if ilvl_el  is not None else '0'
            result[sid] = {'numId': nid, 'ilvl': ilvl}
    except Exception: pass
    return result

def _create_restart_numids(doc: Document, style_numpr: dict) -> dict:
    remapping = {}
    if not style_numpr: return remapping
    try:
        np   = doc.part.numbering_part
        nxml = np._element
        unique_numids = {info['numId'] for info in style_numpr.values()}
        existing      = nxml.findall(f'{_W}num')
        max_id        = max((int(n.get(f'{_W}numId', 0)) for n in existing), default=0)
        for old_nid in unique_numids:
            if old_nid == '0': continue
            old_num_el = None
            for n in existing:
                if n.get(f'{_W}numId') == old_nid:
                    old_num_el = n
                    break
            if old_num_el is None: continue
            max_id += 1
            new_nid    = str(max_id)
            new_num_el = copy.deepcopy(old_num_el)
            new_num_el.set(f'{_W}numId', new_nid)
            for lo in new_num_el.findall(f'{_W}lvlOverride'):
                new_num_el.remove(lo)
            override_el = etree.fromstring(
                f'<w:lvlOverride xmlns:w="{_NS_W}" w:ilvl="0">'
                f'<w:startOverride w:val="1"/>'
                f'</w:lvlOverride>'
            )
            new_num_el.append(override_el)
            nxml.append(new_num_el)
            remapping[old_nid] = new_nid
    except Exception: pass
    return remapping

def _apply_numpr_restart_to_headings(elements: list, style_numpr: dict, remapping: dict) -> None:
    if not remapping: return
    for el in elements:
        if el.tag != f'{_W}p': continue
        sid = _get_style_id(el)
        if sid not in style_numpr: continue
        info    = style_numpr[sid]
        old_nid = info['numId']
        ilvl    = info['ilvl']
        new_nid = remapping.get(old_nid)
        if not new_nid: continue
        pPr = el.find(f'{_W}pPr')
        if pPr is None:
            pPr = etree.SubElement(el, f'{_W}pPr')
            el.insert(0, pPr)
        old_numPr = pPr.find(f'{_W}numPr')
        if old_numPr is not None: pPr.remove(old_numPr)
        pPr.insert(0, etree.fromstring(
            f'<w:numPr xmlns:w="{_NS_W}">'
            f'<w:ilvl w:val="{ilvl}"/>'
            f'<w:numId w:val="{new_nid}"/>'
            f'</w:numPr>'
        ))


# ─────────────────────────────────────────────────────────────────────────────
# HELPER: ELEMEN PEMISAH
# ─────────────────────────────────────────────────────────────────────────────

def _page_break_para() -> etree._Element:
    return etree.fromstring(
        f'<w:p xmlns:w="{_NS_W}">'
        f'<w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>'
        f'<w:r><w:br w:type="page"/></w:r>'
        f'</w:p>'
    )

def _intro_heading() -> etree._Element:
    return etree.fromstring(
        f'<w:p xmlns:w="{_NS_W}">'
        f'<w:pPr>'
        f'<w:jc w:val="center"/>'
        f'<w:spacing w:before="0" w:after="0"/>'
        f'</w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
        f'<w:b/><w:bCs/>'
        f'<w:sz w:val="22"/><w:szCs w:val="22"/>'
        f'</w:rPr>'
        f'<w:t>Introduction</w:t>'
        f'</w:r>'
        f'</w:p>'
    )

def _empty_para() -> etree._Element:
    return etree.fromstring(
        f'<w:p xmlns:w="{_NS_W}">'
        f'<w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>'
        f'</w:p>'
    )


# ─────────────────────────────────────────────────────────────────────────────
# CORE: INSERT ORIGINAL
# ─────────────────────────────────────────────────────────────────────────────

def _insert_original_before_bib(
    translated_doc:    Document,
    orig_body_els:     list,
    progress_callback=None,
) -> None:
    bib_in_orig   = _find_bib_index(orig_body_els)
    content_start = _find_content_start_before_bib(orig_body_els, bib_in_orig)
    content_end   = bib_in_orig if bib_in_orig >= 0 else len(orig_body_els)

    if bib_in_orig < 0 and content_end > content_start:
        if orig_body_els[content_end - 1].tag == f'{_W}sectPr':
            content_end -= 1

    els_to_insert = orig_body_els[content_start:content_end]

    if not els_to_insert:
        _notify(progress_callback, 50, "⚠️ Tidak ada konten asli yang bisa disisipkan.")
        return

    _notify(progress_callback, 10, f"Ditemukan {len(els_to_insert)} elemen.")

    style_numpr = _read_style_numpr(translated_doc)
    _notify(progress_callback, 20, f"Styles: {list(style_numpr.keys())}")

    remapping = _create_restart_numids(translated_doc, style_numpr)
    _notify(progress_callback, 30, f"Remapping: {remapping}")

    new_content_els = [copy.deepcopy(el) for el in els_to_insert]
    _apply_numpr_restart_to_headings(new_content_els, style_numpr, remapping)
    _notify(progress_callback, 40, "Reset heading numbers.")

    trans_body     = translated_doc.element.body
    trans_children = list(trans_body)
    bib_idx_trans  = _find_bib_index(trans_children)

    if bib_idx_trans < 0:
        bib_idx_trans = len(trans_children)
        for i in range(len(trans_children) - 1, -1, -1):
            if trans_children[i].tag == f'{_W}sectPr':
                bib_idx_trans = i
                break
        _notify(progress_callback, 45, "Bib tidak ditemukan → akhir.")
    else:
        _notify(progress_callback, 45, f"Bib posisi [{bib_idx_trans}].")

    new_els = (
        [_page_break_para(), _intro_heading(), _page_break_para()]
        + new_content_els
        + [_empty_para()]
    )

    total = len(new_els)
    for offset, el in enumerate(new_els):
        trans_body.insert(bib_idx_trans + offset, el)
        if progress_callback and offset % 20 == 0:
            pct = 45 + int(offset / max(total, 1) * 55)
            progress_callback(pct, f"Menyisipkan... ({offset}/{total})")

    _notify(progress_callback, 100, f"✅ Selesai ({len(els_to_insert)} item).")


def _notify(cb, pct: int, msg: str) -> None:
    if cb:
        try: cb(pct, msg)
        except Exception: pass


# ─────────────────────────────────────────────────────────────────────────────
# TRANSLATOR WRAPPER — Dengan Custom Dictionary
# ─────────────────────────────────────────────────────────────────────────────

class _Translator:
    def __init__(
        self,
        source: str = 'auto',
        target: str = 'id',
        custom_dict: CustomDictionary | None = None,
    ):
        try:
            from deep_translator import GoogleTranslator
            self._cls        = GoogleTranslator
            self.source      = source
            self.target      = target
            self.custom_dict = custom_dict
        except ImportError:
            raise ImportError("Jalankan: pip install deep-translator")

    def translate_one(self, text: str) -> str:
        t = text.strip()
        if not t or _skip_text(t): return text

        # ── Langkah 1: Terapkan kamus (ganti term → token) ──────────────
        token_map = {}
        if self.custom_dict and len(self.custom_dict) > 0:
            t, token_map = self.custom_dict._apply_pre(t)

        # ── Langkah 2: Jika setelah substitusi teks cukup untuk di-translate
        try:
            result = self._cls(source=self.source, target=self.target).translate(t)
            if not result:
                result = t
        except Exception:
            time.sleep(0.8)
            try:
                result = self._cls(source=self.source, target=self.target).translate(t)
                if not result:
                    result = t
            except Exception:
                result = t

        # ── Langkah 3: Kembalikan token → target ────────────────────────
        if token_map:
            result = self.custom_dict._apply_post(result, token_map)

        return result


# ─────────────────────────────────────────────────────────────────────────────
# CORE: TRANSLATION
# ─────────────────────────────────────────────────────────────────────────────

def _match_capitalization(original: str, translated: str) -> str:
    """
    Samakan pola kapitalisasi terjemahan dengan teks asli.

    Aturan:
      1. Original ALL CAPS (≥ 80% huruf besar)  → translated.upper()
      2. Original diawali huruf kapital           → Sentence case (hanya huruf pertama)
      3. Original lowercase semua                → biarkan hasil translate apa adanya
    """
    orig = original.strip()
    tran = translated.strip()
    if not orig or not tran:
        return translated

    letters = [c for c in orig if c.isalpha()]
    if not letters:
        return translated

    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)

    # ALL CAPS (mis: "SCOPE", "NORMATIVE REFERENCES")
    if upper_ratio >= 0.8:
        return tran.upper()

    # Sentence case — hanya huruf pertama kapital jika original diawali kapital
    # Ini yang benar untuk heading ISO: "Scope" → "Ruang lingkup"
    if orig[0].isupper():
        return tran[0].upper() + tran[1:] if len(tran) > 1 else tran.upper()

    # Lowercase
    return tran


def _translate_para(para, tr, past_bibliography: bool = False) -> None:
    if _skip_paragraph(para, past_bibliography): return

    text_runs = [(i, r) for i, r in enumerate(para.runs)
                 if r.text and r.text.strip()]
    if not text_runs: return

    combined = ''.join(r.text for _, r in text_runs)
    if _skip_text(combined): return

    translated = tr.translate_one(combined.strip())
    time.sleep(_TRANSLATE_DELAY)
    if not translated or translated == combined: return

    # Samakan pola kapital dengan teks asli
    translated = _match_capitalization(combined.strip(), translated)

    _, first_run = text_runs[0]
    first_run.text = translated
    for _, run in text_runs[1:]:
        run.text = ''

def _translate_table(table, tr) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _translate_para(para, tr)

def _translate_hf(hf_part, tr) -> None:
    if hf_part is None: return
    try:
        for para in hf_part.paragraphs:
            if _RE_COPYRIGHT.search(para.text or ''): continue
            _translate_para(para, tr)
        for table in hf_part.tables:
            _translate_table(table, tr)
    except Exception: pass


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ENGINE CLASS
# ─────────────────────────────────────────────────────────────────────────────

class DocxFinalTranslatorEngine:
    """
    Engine 9:
      - Semua fitur Engine 8 (Annex, em-dash bullet, Introduction page).
      - **BARU**: Custom Dictionary dengan prioritas di atas Google Translate.
        Dukung input manual, CSV, dan Excel.

    Contoh:
        from engine9 import DocxFinalTranslatorEngine, CustomDictionary

        d = CustomDictionary()
        d.load_defaults()
        d.load_from_csv("kamus_proyek.csv")
        d.add_term("bearing pad", "bantalan elastomer")

        engine = DocxFinalTranslatorEngine(custom_dict=d)
        ok, msg = engine.translate("input.docx", "output.docx")
    """

    def __init__(
        self,
        source_lang:  str = 'auto',
        target_lang:  str = 'id',
        custom_dict:  CustomDictionary | None = None,
    ):
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.custom_dict = custom_dict  # bisa None → tidak pakai kamus

    # ── AKSES CEPAT KE KAMUS ──────────────────────────────────────────────

    def set_dictionary(self, d: CustomDictionary) -> None:
        """Set atau ganti kamus."""
        self.custom_dict = d

    def get_dictionary(self) -> CustomDictionary:
        """Ambil kamus; buat baru jika belum ada."""
        if self.custom_dict is None:
            self.custom_dict = CustomDictionary()
        return self.custom_dict

    # ── TRANSLATE ─────────────────────────────────────────────────────────

    def translate(
        self,
        input_docx:        str,
        output_docx:       str,
        progress_callback=None,
        translate_headers: bool = False,
    ) -> tuple[bool, str]:
        try:
            dict_info = (
                f"{len(self.custom_dict)} istilah kamus"
                if self.custom_dict else "tanpa kamus"
            )
            _notify(progress_callback, 2, f"Snapshot dokumen... ({dict_info})")
            doc_orig      = Document(input_docx)
            orig_body_els = [copy.deepcopy(el) for el in doc_orig.element.body]
            del doc_orig

            _notify(progress_callback, 5, "Init translator...")
            tr  = _Translator(
                source=self.source_lang,
                target=self.target_lang,
                custom_dict=self.custom_dict,
            )
            doc = Document(input_docx)

            body     = doc.element.body
            para_map = {p._element: p for p in doc.paragraphs}
            tbl_map  = {t._element: t for t in doc.tables}

            items    = []
            sec_brks = 0
            COVER_END = 1

            for child in body:
                if child in para_map:
                    in_cover = (sec_brks < COVER_END)
                    items.append(('para', para_map[child], in_cover))
                    if _has_inline_sectpr(para_map[child]):
                        sec_brks += 1
                elif child in tbl_map:
                    in_cover = (sec_brks < COVER_END)
                    items.append(('table', tbl_map[child], in_cover))

            total             = len(items)
            done              = 0
            past_bibliography = False

            for kind, obj, in_cover in items:
                done += 1
                pct = 5 + int(done / max(total, 1) * 60)

                if kind == 'para':
                    para           = obj
                    is_bib_heading = _is_biblio_title_para(para)
                    is_annex       = _get_para_style_id(para) in _ANNEX_STYLE_IDS

                    if in_cover and _all_runs_italic(para):
                        _notify(progress_callback, pct, "[Cover-italic] skip")

                    elif is_bib_heading:
                        _translate_para(para, tr, past_bibliography=False)
                        past_bibliography = True
                        _notify(progress_callback, pct, "[Bib-heading]")

                    elif is_annex and not past_bibliography:
                        _translate_para(para, tr, past_bibliography=False)
                        _fix_annex_style_para(para)
                        _notify(progress_callback, pct, "[ANNEX] fixed")

                    elif not _skip_paragraph(para, past_bibliography):
                        _translate_para(para, tr, past_bibliography=False)

                elif kind == 'table':
                    if not past_bibliography:
                        _translate_table(obj, tr)

            _notify(progress_callback, 66, "Formatting em-dash bullets...")
            _convert_emdash_to_bullets(doc)

            if translate_headers:
                _notify(progress_callback, 70, "Translating headers/footers...")
                for section in doc.sections:
                    for hf in [
                        section.header,            section.footer,
                        section.even_page_header,  section.even_page_footer,
                        section.first_page_header, section.first_page_footer,
                    ]:
                        _translate_hf(hf, tr)

            _notify(progress_callback, 75, "Inserting original content...")

            def _insert_cb(pct_inner, msg):
                _notify(progress_callback, 75 + int(pct_inner * 0.20), msg)

            _insert_original_before_bib(
                translated_doc=doc,
                orig_body_els=orig_body_els,
                progress_callback=_insert_cb,
            )

            _notify(progress_callback, 97, "Saving...")
            doc.save(output_docx)
            _notify(progress_callback, 100, "✅ Done!")
            return True, output_docx

        except ImportError as e:
            return False, f"Dependensi tidak ditemukan: {e}\nJalankan: pip install deep-translator pandas openpyxl"
        except Exception as e:
            return False, f"Engine9 Error: {str(e)}\n{traceback.format_exc()}"